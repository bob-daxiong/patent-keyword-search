import os
import docx


def parse_docx(file_path: str) -> str:
    doc = docx.Document(file_path)
    parts = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    for table in doc.tables:
        for row in table.rows:
            cells = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    cells.append(cell_text)
            if cells:
                parts.append(' '.join(cells))

    for section in doc.sections:
        for para in section.header.paragraphs:
            text = para.text.strip()
            if text:
                parts.append(text)
        for para in section.footer.paragraphs:
            text = para.text.strip()
            if text:
                parts.append(text)

    return '\n'.join(parts)


def cleanup_file(file_path: str) -> None:
    if os.path.exists(file_path):
        os.remove(file_path)
